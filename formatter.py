from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# =====================================================
# MAIN FUNCTION
# =====================================================

def format_document(input_path, output_path):
    doc = Document(input_path)
    set_default_font(doc)
    format_paragraphs(doc)
    format_tables(doc)
    add_header_footer(doc)
    add_page_number(doc)
    doc.save(output_path)


# =====================================================
# GLOBAL DEFAULT FONT
# =====================================================

def set_default_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# =====================================================
# TEXT HELPERS
# =====================================================

def sentence_case(text):
    if not text:
        return text
    text = text.strip()
    return text[0].upper() + text[1:].lower()


def is_number_only(text):
    return bool(re.fullmatch(r'[\d.\-/:]+', text.strip()))


# =====================================================
# PARAGRAPH FORMATTING
# =====================================================

def format_paragraphs(doc):
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        lower = text.lower()

        # PAGE 1 MAIN TITLE
        if i < 3 and len(text) > 15:
            apply_main_title(para, text)

        # CURRICULUM / DEGREE TITLES (PAGE 2)
        elif "curriculum" in lower or "b.sc" in lower or "m.sc" in lower:
            apply_blue_center_heading(para, text, 14)

        # PROGRAM INFO
        elif "programme" in lower or "applicable" in lower:
            apply_center_bold(para, text, 14)

        # SEMESTER HEADING (PAGE 3)
        elif "semester" in lower:
            apply_blue_center_heading(para, text, 14)

        # SECTION HEADINGS
        elif any(k in lower for k in ["eligibility", "objectives", "outcomes", "references", "text books", "web resources"]):
            apply_left_bold(para, text, 12)

        # BULLETS
        elif text.startswith(("•", "-", "1.", "2.", "3.")):
            apply_body(para, text, single=True)

        # REFERENCES (APA style lines)
        elif "(" in text and ")" in text and "," in text:
            apply_body(para, text, single=True)

        # DEFAULT BODY
        else:
            apply_body(para, text)


# ================= STYLE APPLIERS =================

def clear_and_write(para, text):
    para.clear()
    return para.add_run(sentence_case(text))


def apply_main_title(para, text):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = clear_and_write(para, text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)


def apply_blue_center_heading(para, text, size):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = clear_and_write(para, text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)


def apply_center_bold(para, text, size):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = clear_and_write(para, text)
    run.font.size = Pt(size)
    run.font.bold = True


def apply_left_bold(para, text, size):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = clear_and_write(para, text)
    run.font.size = Pt(size)
    run.font.bold = True


def apply_body(para, text, single=False):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE if single else WD_LINE_SPACING.ONE_POINT_FIVE
    run = clear_and_write(para, text)
    run.font.size = Pt(12)


# =====================================================
# TABLE FORMATTING
# =====================================================

def format_tables(doc):
    for table in doc.tables:
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    para.clear()
                    run = para.add_run(sentence_case(text))
                    run.font.size = Pt(12)

                    if r == 0:  # header row
                        run.font.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade_cell(cell)
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_number_only(text) else WD_ALIGN_PARAGRAPH.LEFT


def shade_cell(cell):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), "D9D9D9")
    tcPr.append(shd)


# =====================================================
# HEADER / FOOTER
# =====================================================

def add_header_footer(doc):
    section = doc.sections[0]
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)

    header = section.header.paragraphs[0]
    header.text = "KG College of Arts and Science (Autonomous) – 2024 Batch"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(11)
    header.runs[0].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.text = "Department of Computer Science"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(11)
    footer.runs[0].font.bold = True


# =====================================================
# PAGE NUMBER
# =====================================================

def add_page_number(doc):
    footer = doc.sections[0].footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
